"""Write test-case rows to xlsx / csv / json / txt / docx / pdf.

xlsx output includes a Coverage sheet (endpoints × categories matrix) and
preserves rows previously marked Reviewed/Approved unless --force is set.
"""
import argparse
import csv
import datetime as dt
import json
import logging
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any

import yaml

LOG = logging.getLogger(__name__)

_SKILL_ROOT = Path(__file__).parent.parent
_COLUMNS_CFG = _SKILL_ROOT / "templates" / "testcase_columns.yaml"
_OUTPUT_DIR = _SKILL_ROOT / "generated_spec_sheets"


def _load_columns() -> list[dict[str, Any]]:
    data = yaml.safe_load(_COLUMNS_CFG.read_text())
    return list(data["columns"])


def _output_path(source: str, component: str, fmt: str) -> Path:
    ts = dt.datetime.now(dt.timezone.utc).strftime("%Y%m%d_%H%M%SZ")
    return _OUTPUT_DIR / f"testcases_{source}_{component}_{ts}.{fmt}"


def _stringify(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    if value is None:
        return ""
    return str(value)


def _merge_with_previous(rows: list[dict[str, Any]], previous: Path | None,
                         force: bool) -> tuple[list[dict[str, Any]], int]:
    if previous is None or not previous.exists():
        return rows, 0
    try:
        from openpyxl import load_workbook
        wb = load_workbook(previous)
        ws = wb["TestCases"]
        headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        prev_rows: dict[str, dict[str, Any]] = {}
        for row in ws.iter_rows(min_row=2, values_only=True):
            entry = dict(zip(headers, row))
            if entry.get("tc_id"):
                prev_rows[entry["tc_id"]] = entry
    except Exception as exc:
        LOG.info("Could not load previous xlsx (%s) — proceeding without merge", exc)
        return rows, 0

    preserved = 0
    by_id = {r["tc_id"]: r for r in rows}
    for tc_id, prev in prev_rows.items():
        status = (prev.get("review_status") or "Draft").strip()
        if status in ("Reviewed", "Approved") and not force and tc_id in by_id:
            by_id[tc_id] = prev
            preserved += 1
    return list(by_id.values()), preserved


def _build_coverage(rows: list[dict[str, Any]]) -> dict[tuple[str, str], int]:
    matrix: dict[tuple[str, str], int] = defaultdict(int)
    for r in rows:
        matrix[(f"{r['method']} {r['endpoint_url']}", r["category"])] += 1
    return matrix


def write_xlsx(rows: list[dict[str, Any]], columns: list[dict[str, Any]], path: Path,
               previous: Path | None) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    if ws is None:
        raise RuntimeError("openpyxl returned no active sheet")
    ws.title = "TestCases"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="305496")
    col_names = [c["name"] for c in columns]
    ws.append(col_names)
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
    ws.freeze_panes = "A2"

    category_colors = {
        "positive": "C6EFCE", "negative_validation": "FFEB9C",
        "negative_auth": "FFC7CE", "negative_not_found": "FFD9B3",
        "negative_conflict": "F4B084", "negative_rate": "E4DFEC",
        "schema": "BDD7EE", "boundary": "FFE699", "i18n": "D9E1F2",
        "idempotency": "EDEDED", "cross_env": "B4C7E7", "performance": "C9C9C9",
    }
    cat_col_idx = col_names.index("category") + 1

    for row in rows:
        ws.append([_stringify(row.get(name, "")) for name in col_names])
        cell = ws.cell(row=ws.max_row, column=cat_col_idx)
        color = category_colors.get(str(cell.value), None)
        if color:
            cell.fill = PatternFill("solid", fgColor=color)

    for i, name in enumerate(col_names, start=1):
        ws.column_dimensions[get_column_letter(i)].width = min(max(len(name) + 2, 14), 50)

    # Coverage sheet
    cov_ws = wb.create_sheet("Coverage")
    matrix = _build_coverage(rows)
    endpoints = sorted({k[0] for k in matrix})
    categories = sorted({k[1] for k in matrix})
    cov_ws.append(["endpoint"] + categories)
    for cell in cov_ws[1]:
        cell.font = header_font
        cell.fill = header_fill
    for ep in endpoints:
        cov_ws.append([ep] + [matrix.get((ep, cat), 0) for cat in categories])
    cov_ws.freeze_panes = "B2"

    # Delta sheet (if comparing)
    if previous and previous.exists():
        try:
            from openpyxl import load_workbook
            wb_prev = load_workbook(previous)
            ws_prev = wb_prev["TestCases"]
            headers = [c.value for c in next(ws_prev.iter_rows(min_row=1, max_row=1))]
            prev_ids = set()
            for r in ws_prev.iter_rows(min_row=2, values_only=True):
                entry = dict(zip(headers, r))
                if entry.get("tc_id"):
                    prev_ids.add(entry["tc_id"])
            new_ids = {r["tc_id"] for r in rows}
            delta_ws = wb.create_sheet("Delta")
            delta_ws.append(["tc_id", "change"])
            for tid in sorted(new_ids - prev_ids):
                delta_ws.append([tid, "added"])
            for tid in sorted(prev_ids - new_ids):
                delta_ws.append([tid, "removed"])
        except Exception as exc:
            LOG.info("Delta sheet skipped (%s)", exc)

    wb.save(path)


def write_csv(rows: list[dict[str, Any]], columns: list[dict[str, Any]], path: Path) -> None:
    col_names = [c["name"] for c in columns]
    with open(path, "w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(col_names)
        for row in rows:
            writer.writerow([_stringify(row.get(name, "")) for name in col_names])


def write_json(rows: list[dict[str, Any]], path: Path) -> None:
    path.write_text(json.dumps(rows, indent=2, ensure_ascii=False))


def write_txt(rows: list[dict[str, Any]], columns: list[dict[str, Any]], path: Path) -> None:
    col_names = [c["name"] for c in columns]
    lines: list[str] = []
    for row in rows:
        lines.append("=" * 72)
        for name in col_names:
            lines.append(f"{name}: {_stringify(row.get(name, ''))}")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_docx(rows: list[dict[str, Any]], columns: list[dict[str, Any]], path: Path) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading("Test Case Specification", level=1)
    col_names = [c["name"] for c in columns]
    table = doc.add_table(rows=1, cols=len(col_names))
    table.style = "Light Grid Accent 1"
    for i, name in enumerate(col_names):
        table.rows[0].cells[i].text = name
    for row in rows:
        cells = table.add_row().cells
        for i, name in enumerate(col_names):
            cells[i].text = _stringify(row.get(name, ""))
    doc.save(str(path))


def write_pdf(rows: list[dict[str, Any]], columns: list[dict[str, Any]], path: Path) -> None:
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    col_names = [c["name"] for c in columns]
    data = [col_names] + [[_stringify(row.get(n, "")) for n in col_names] for row in rows]
    doc = SimpleDocTemplate(str(path), pagesize=landscape(letter))
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#305496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
    ]))
    doc.build([table])


_WRITERS = {
    "xlsx": "xlsx", "csv": "csv", "json": "json",
    "txt": "txt", "docx": "docx", "pdf": "pdf",
}


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--rows", required=True, help="Path to rows JSON")
    ap.add_argument("--format", default="xlsx", choices=list(_WRITERS))
    ap.add_argument("--source", default="manual")
    ap.add_argument("--component", default="mixed")
    ap.add_argument("--against", default=None, help="Previous xlsx path for diff/merge")
    ap.add_argument("--force", action="store_true", help="Overwrite Reviewed/Approved rows")
    args = ap.parse_args()

    rows_path = Path(args.rows)
    if not rows_path.exists():
        print(f"[write_output] rows file not found: {rows_path}", file=sys.stderr)
        sys.exit(2)
    rows: list[dict[str, Any]] = json.loads(rows_path.read_text())

    previous = Path(args.against) if args.against else None
    rows, preserved = _merge_with_previous(rows, previous, args.force)

    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = _output_path(args.source, args.component, args.format)
    columns = _load_columns()

    if args.format == "xlsx":
        write_xlsx(rows, columns, out, previous)
    elif args.format == "csv":
        write_csv(rows, columns, out)
    elif args.format == "json":
        write_json(rows, out)
    elif args.format == "txt":
        write_txt(rows, columns, out)
    elif args.format == "docx":
        write_docx(rows, columns, out)
    elif args.format == "pdf":
        write_pdf(rows, columns, out)

    summary = Counter(r.get("category", "?") for r in rows)
    priority = Counter(r.get("priority", "?") for r in rows)
    print(f"Wrote {out}")
    print(f"Rows: {len(rows)}  Preserved (Reviewed/Approved): {preserved}")
    print(f"By category: {dict(summary)}")
    print(f"By priority: {dict(priority)}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    main()
